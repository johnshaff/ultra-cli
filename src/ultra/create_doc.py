from docx import Document
from docx.shared import Pt, RGBColor
import os
import json


def write_styled_docx(json_file: str):
    # Load metadata from the JSON file
    with open(json_file, "r") as jf:
        data = json.load(jf)

    # Create document and set output path using the JSON id field
    doc = Document()
    output_path = f"transcript/{data['id']}.docx"

    # Make headings black and set font size for Heading 1
    heading1_style = doc.styles["Heading 1"]
    heading1_style.font.color.rgb = RGBColor(0, 0, 0)
    heading1_style.font.size = Pt(20)
    heading2_style = doc.styles["Heading 2"]
    heading2_style.font.color.rgb = RGBColor(0, 0, 0)

    # 1) 🎬 Video Title (Heading 1) - remains hardcoded as requested
    doc.add_heading("🎬 Video Title", level=1)

    # 2) Quoted subtitle (Heading 2) - now taken from JSON field "title"
    doc.add_heading(f'"{data["title"]}"', level=2)

    # 3) Separator "⸻"
    sep_par_1 = doc.add_paragraph()
    sep_par_1.paragraph_format.space_before = Pt(0)
    sep_par_1.paragraph_format.space_after = Pt(0)
    sep_par_1.add_run("----------------------------------------------------------------------------------------------------------------------").font.color.rgb = RGBColor(0, 0, 0)

    # 4) 🎥 Video Snapshot (Heading 2)
    doc.add_heading("🎥 Video Snapshot", level=2)

    # 5) Single paragraph with bold labels, with metadata extracted from JSON
    snap_par = doc.add_paragraph()
    snap_par.add_run("Duration: ").bold = True
    snap_par.add_run(f"{data['duration']} | ")

    snap_par.add_run("Uploaded: ").bold = True
    snap_par.add_run(f"{data['upload_date']} | ")

    snap_par.add_run("Views: ").bold = True
    snap_par.add_run(f"{data['view_count']} | ")

    snap_par.add_run("Likes: ").bold = True
    snap_par.add_run(f"{data['like_count']} | ")

    snap_par.add_run("Comments: ").bold = True
    snap_par.add_run(f"{data['comment_count']} | ")

    snap_par.add_run("Channel: ").bold = True
    snap_par.add_run(f"{data['uploader']} ({data['uploader_id']}) | ")

    snap_par.add_run("Category: ").bold = True
    snap_par.add_run(data["categories"])

    # 6) Description heading
    doc.add_heading("📝 Description:", level=2)

    # 7) Description paragraph from JSON field "description"
    doc.add_paragraph(data["description"])

    # 8) Second separator
    sep_par_2 = doc.add_paragraph()
    sep_par_2.paragraph_format.space_before = Pt(0)
    sep_par_2.paragraph_format.space_after = Pt(0)
    sep_par_2.add_run("----------------------------------------------------------------------------------------------------------------------").font.color.rgb = RGBColor(0, 0, 0)

    # 9) Document Transcript heading
    doc.add_heading("📜 Video Transcript", level=2)

    # 10) Transcript content from file using the id from JSON
    transcript_file_path = f"transcript/{data['id']}-final.txt"
    #transcript_file_path = f"transcript/KZcXSG5r5q4-final.txt"
    with open(transcript_file_path, "r") as file:
        transcript_content = file.read()
        doc.add_paragraph(transcript_content)

    # Save and open the document
    doc.save(output_path)
    print(f"Document has been created: {output_path}")
    os.system(f"open {output_path}")


if __name__ == "__main__":
    write_styled_docx()